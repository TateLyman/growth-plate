#!/usr/bin/env python3
"""
HOLDINGS - what this atlas physically has on disk, and whether it has read it.

WHY THIS EXISTS. Four corrections in three rounds have the same shape:

  CORR-232  recommended an archive request without querying the archive's index; and the number
            being hunted was inside a reference already in the bibliography
  CORR-234  a paper was filed unreadable in round 45 on one failed retrieval, never retested, and
            spent twenty rounds being reconstructed from case reports
  CORR-235  a node's headline patient was assigned the wrong drug, because the drug was in a table
            of a document already held and read for a different purpose
  round 230 an ASK LIST was issued to the user containing the ASCO abstract 10007 - which was
            sitting in atlas/data/round227_supplied/ having been supplied three rounds earlier

That is not four mistakes. It is one missing function: THE ATLAS COULD NOT ANSWER THE QUESTION
"DO I ALREADY HAVE THIS?" This file is that function.

  holdings.py                      rebuild the index and print the summary
  holdings.py --find "bone age"    full-text search EVERY held document
  holdings.py --have 37158537      do I hold this PMID / DOI / NCT / ref_id? exit 0 if yes
  holdings.py --check-asks         scan atlas/audit/ASK_LIST*.md and FAIL on anything already held
                                   as a DOCUMENT or already answered in a ref's one_line_finding
  holdings.py --unread             bibliography refs with no local artefact, ranked by citation count
  holdings.py --orphans            held documents that NO bibliography ref names - the root cause

THE RULE THIS ENFORCES: run --check-asks before sending any request to the user, and --have before
declaring anything unobtainable. An ask list that names a document already on disk is a defect.

Writes atlas/sources/holdings.json and atlas/sources/HOLDINGS.md.
"""

import argparse
import json
import os
import re
import sys

import yaml

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.normpath(os.path.join(HERE, ".."))
DATA = os.path.join(ROOT, "data")
BIB = os.path.join(ROOT, "sources", "bibliography.yaml")
INDEX_JSON = os.path.join(ROOT, "sources", "holdings.json")
INDEX_MD = os.path.join(ROOT, "sources", "HOLDINGS.md")

TEXTLIKE = {".txt", ".md", ".json", ".csv", ".tsv", ".xml", ".yaml", ".yml", ".html", ".htm"}
BINARY_DOC = {".pdf", ".docx", ".xlsx", ".pptx", ".png", ".jpg", ".jpeg", ".gz", ".zip"}

PMID_RE = re.compile(r"\b(?:PMID[:\s]*)?(\d{7,8})\b")
PMCID_RE = re.compile(r"\bPMC(\d{6,9})\b")
DOI_RE = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Za-z0-9]+\b")
NCT_RE = re.compile(r"\bNCT\d{8}\b")


def read_text(path, limit=400_000):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in TEXTLIKE:
            with open(path, errors="replace") as fh:
                return fh.read(limit)
        if ext == ".pdf":
            import pymupdf
            doc = pymupdf.open(path)
            return "\n".join(p.get_text() for p in doc)[:limit]
        if ext == ".docx":
            from docx import Document
            d = Document(path)
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for r in t.rows:
                    parts.append(" | ".join(c.text for c in r.cells))
            return "\n".join(parts)[:limit]
    except Exception as e:
        return f"<<unreadable: {e}>>"
    return ""


def build():
    bib = yaml.safe_load(open(BIB))["refs"]
    by_pmid, by_doi = {}, {}
    for k, v in bib.items():
        if v.get("pmid"):
            by_pmid[str(v["pmid"])] = k
        if v.get("doi"):
            by_doi[str(v["doi"]).lower()] = k

    items = []
    for dirpath, _dirs, files in os.walk(DATA):
        for fn in sorted(files):
            path = os.path.join(dirpath, fn)
            rel = os.path.relpath(path, ROOT)
            ext = os.path.splitext(fn)[1].lower()
            size = os.path.getsize(path)
            body = read_text(path) if (ext in TEXTLIKE or ext in {".pdf", ".docx"}) else ""
            pmids = sorted(set(PMID_RE.findall(body)) & set(by_pmid)) if body else []
            dois = sorted({d.lower().rstrip(".,);") for d in DOI_RE.findall(body)}) if body else []
            ncts = sorted(set(NCT_RE.findall(body))) if body else []
            refs = sorted({by_pmid[p] for p in pmids} |
                          {by_doi[d] for d in dois if d in by_doi})
            # a ref_id appearing in the filename is a strong, deliberate signal
            stem = re.sub(r"[^a-z0-9]", "", fn.lower())
            named = sorted({k for k in bib if len(k) > 6 and re.sub(r"[^a-z0-9]", "", k) in stem})
            items.append({
                "path": rel, "bytes": size, "ext": ext,
                "chars": len(body),
                "refs_cited_inside": refs,
                "ref_in_filename": named,
                "ncts": ncts[:12],
                "snippet": re.sub(r"\s+", " ", body[:300]).strip(),
            })

    held_refs = sorted({r for i in items for r in i["ref_in_filename"]})
    idx = {"items": items, "held_refs_by_filename": held_refs,
           "n_files": len(items), "n_bytes": sum(i["bytes"] for i in items)}
    json.dump(idx, open(INDEX_JSON, "w"), indent=1)

    with open(INDEX_MD, "w") as fh:
        fh.write("# HOLDINGS — every document this atlas physically has\n\n")
        fh.write("Generated by `atlas/tools/holdings.py`. **Do not hand-edit; rerun the tool.**\n\n")
        fh.write("Before asking anyone for a document, run `holdings.py --check-asks`. Before calling\n"
                 "anything unobtainable, run `holdings.py --have <pmid|doi|NCT|ref_id>`.\n\n")
        fh.write(f"- files: **{len(items)}**\n- bytes: **{sum(i['bytes'] for i in items):,}**\n")
        fh.write(f"- bibliography refs with a file named after them: **{len(held_refs)}**\n\n")
        fh.write("## Refs held as a named local document\n\n")
        fh.write(", ".join(f"`{r}`" for r in held_refs) + "\n\n")
        fh.write("## Every file\n\n| path | size | chars | refs inside | trials |\n|---|---|---|---|---|\n")
        for i in sorted(items, key=lambda x: x["path"]):
            fh.write(f"| `{i['path']}` | {i['bytes']:,} | {i['chars']:,} | "
                     f"{', '.join(i['refs_cited_inside'][:6]) or '-'} | "
                     f"{', '.join(i['ncts'][:4]) or '-'} |\n")
    return idx


def load():
    if not os.path.exists(INDEX_JSON):
        return build()
    return json.load(open(INDEX_JSON))


def cmd_find(term):
    idx = load()
    pat = re.compile(term, re.I)
    hits = 0
    for i in idx["items"]:
        path = os.path.join(ROOT, i["path"])
        body = read_text(path)
        if not body:
            continue
        for m in pat.finditer(body):
            hits += 1
            ctx = re.sub(r"\s+", " ", body[max(0, m.start() - 160):m.start() + 200])
            print(f"\n{i['path']}\n    ...{ctx}...")
            break
    print(f"\n{hits} of {idx['n_files']} held files contain /{term}/")
    return 0 if hits else 1


STOP = {"the", "of", "in", "with", "from", "and", "for", "a", "an", "to", "on", "results",
        "study", "patients", "trial", "using", "analysis", "report", "case"}


def title_probe(token):
    """If the token names a bibliography ref, return a distinctive phrase from its title.

    Identifier matching alone is not enough and that is not a theoretical worry: a conference
    abstract PDF does not print its own DOI, which is exactly how the ASCO 10007 abstract sat
    in atlas/data/round227_supplied/ while being requested from the user by DOI.
    """
    try:
        bib = yaml.safe_load(open(BIB))["refs"]
    except Exception:
        return None, None
    t = token.strip().lower()
    for k, v in bib.items():
        if (k.lower() == t or str(v.get("pmid")) == t or str(v.get("doi", "")).lower() == t):
            title = str(v.get("title", ""))
            words = [w for w in re.findall(r"[A-Za-z0-9-]{3,}", title) if w.lower() not in STOP]
            return k, " ".join(words[:5]) if len(words) >= 3 else None
    return None, None


def cmd_have(token):
    idx = load()
    t = token.strip().lower()
    ref_id, phrase = title_probe(token)
    if phrase:
        pat = re.compile(r"\W+".join(re.escape(w) for w in phrase.split()), re.I)
        thits = []
        for i in idx["items"]:
            body = read_text(os.path.join(ROOT, i["path"]))
            if body and pat.search(body):
                thits.append(i["path"])
        if thits:
            print(f"HELD BY TITLE - '{token}' resolves to ref '{ref_id}', whose title phrase")
            print(f"  \"{phrase}\" appears in {len(thits)} local file(s):")
            for f in thits[:20]:
                print("   ", f)
            print("\nDO NOT ask anyone for this. Read what is already on disk.")
            return 0
    # STRONG means a file plausibly IS the document. WEAK means it is merely mentioned - which is
    # what an entry in somebody else's reference list looks like, and reporting that as HELD would
    # cause the opposite failure to CORR-238: not asking for something genuinely needed.
    strong, weak = [], []
    for i in idx["items"]:
        in_name = t in os.path.basename(i["path"]).lower()
        by_ref = t in {r.lower() for r in i["ref_in_filename"]}
        if in_name or by_ref:
            strong.append((i["path"], -1))
            continue
        body = read_text(os.path.join(ROOT, i["path"]))
        if not body:
            continue
        n = body.lower().count(t)
        if not n:
            continue
        manifest = os.path.basename(i["path"]).lower().startswith(("manifest", "search"))
        (strong if (n >= 3 and i["chars"] > 4000 and not manifest) else weak).append((i["path"], n))
    if strong:
        print(f"HELD - '{token}' looks present as a document in {len(strong)} local file(s):")
        for f, n in strong[:20]:
            print(f"    {f}" + (f"  ({n} mentions)" if n > 0 else "  (named for it)"))
        print("\nDO NOT ask anyone for this. Read what is already on disk.")
        return 0
    if weak:
        print(f"NOT HELD as a document - '{token}' is only MENTIONED in {len(weak)} file(s), which is what")
        print("a citation in somebody else's reference list looks like. Asking for it is legitimate:")
        for f, n in weak[:8]:
            print(f"    {f}  ({n} mention{'s' if n > 1 else ''})")
        return 1
    print(f"not held - '{token}' is not in any file under atlas/data/")
    return 1


def bib_answer_probe(text):
    """Does the BIBLIOGRAPHY already answer what this ask list is asking for?

    holdings.py originally checked only atlas/data/ - whether a DOCUMENT was on disk. That is not
    the whole memory. A ref can carry its key findings in `one_line_finding` without the PDF ever
    having been held, and round 235 found exactly that: the round-233 ask list requested the PROPEL3
    supplement while `propel3_2026` sat in the bibliography with the randomised effect size, the
    bone-age statement and the FGFR1 statement written into its one_line_finding.

    So: pull the distinctive content words out of each ask, and report any ref whose recorded
    finding covers several of them.

    KNOWN LIMITATION, stated because a check whose limits are hidden is worse than no check. This
    pass is tuned for PRECISION over recall - a word must appear in five or fewer findings across
    the whole bibliography to count, and three must match. A noisy version of this check was built
    first and discarded: it fired on "data", "growth", "label" and "outcome", and a check that
    cries wolf gets ignored, which is the exact failure it exists to prevent. THE CONSEQUENCE IS
    THAT IT WILL MISS CASES. It does not catch `propel3_2026` from a heading reading "PROPEL 3
    supplementary appendix", because the shared tokens are too common. THE STANDING RULE THEREFORE
    REMAINS MANUAL AND COMES FIRST: before asking any human for a document, grep the bibliography
    for the study name. This tool narrows that duty; it does not discharge it.
    """
    bib = yaml.safe_load(open(BIB))["refs"]
    findings = {k: str(v.get("one_line_finding") or "").lower() for k, v in bib.items()}
    findings = {k: f for k, f in findings.items() if f}
    if not findings:
        return []
    # A word only counts if it is RARE across the corpus of recorded findings. Without this the
    # probe fires on "data", "growth" and "review" and is immediately worthless - a noisy check
    # gets ignored, which defeats the point of having one.
    import collections
    df = collections.Counter()
    for f in findings.values():
        df.update({w for w in re.findall(r"[a-z][a-z0-9-]{4,}", f)})
    rare_max = 5          # absolute: a word in >5 findings is not distinctive enough to match on
    out = []
    for h in re.findall(r"^#{2,3}\s*(.+)$", text, re.M):
        if not re.match(r"^\**\d", h.strip()):     # only numbered ASKS, not status or table headings
            continue
        words = {w.lower() for w in re.findall(r"[A-Za-z][A-Za-z0-9-]{4,}", h)} - STOP
        words = {w for w in words if 0 < df.get(w, 0) <= rare_max}
        if len(words) < 2:
            continue
        for k, f in findings.items():
            hit = {w for w in words if w in f}
            if len(hit) >= 3:
                out.append((h.strip()[:70], k, sorted(hit)[:6], f[:220]))
    return out


def cmd_check_asks():
    idx = load()
    import glob
    problems = 0
    for md in sorted(glob.glob(os.path.join(ROOT, "audit", "ASK_LIST*.md"))):
        text = open(md).read()
        tokens = set(PMID_RE.findall(text)) | {d.lower().rstrip(".,);") for d in DOI_RE.findall(text)} \
                 | set(NCT_RE.findall(text))
        for tok in sorted(tokens):
            strong, weak = [], []
            for i in idx["items"]:
                body = read_text(os.path.join(ROOT, i["path"]))
                if not body:
                    continue
                n = body.lower().count(tok.lower())
                if not n:
                    continue
                # STRONG = this file plausibly IS the document: named for it, or it is a
                # substantial text that refers to it repeatedly. WEAK = an incidental mention,
                # which a manifest or a search log will always have and which proves nothing.
                in_name = tok.lower() in os.path.basename(i["path"]).lower()
                manifest = os.path.basename(i["path"]).lower().startswith(("manifest", "search"))
                if in_name or (n >= 3 and i["chars"] > 4000 and not manifest):
                    strong.append((i["path"], n))
                else:
                    weak.append((i["path"], n))
            if strong:
                problems += 1
                print(f"*** {os.path.relpath(md, ROOT)} asks for {tok} - LIKELY ALREADY HELD:")
                for h, n in strong[:4]:
                    print(f"         {h}  ({n} mentions)")
            elif weak:
                print(f"  (weak) {tok} mentioned in {len(weak)} file(s) but no file looks like the "
                      f"document itself - e.g. {weak[0][0]}")
        # and the second memory: does a bibliography ref already record the answer?
        for head, ref, hit, finding in bib_answer_probe(text):
            problems += 1
            print(f"*** {os.path.relpath(md, ROOT)} asks under \"{head}\"")
            print(f"    but ref '{ref}' ALREADY RECORDS A FINDING covering {hit}:")
            print(f"      {finding}")
    if problems:
        print(f"\n{problems} ask(s) are already covered by a held document or a recorded finding. "
              f"CHECK EACH BEFORE SENDING THE LIST.")
        return 1
    print("no ask is covered by a held document or a recorded bibliography finding")
    return 0


def cmd_orphans():
    """Held documents that no bibliography ref points at.

    THIS IS THE ROOT CAUSE OF THE FAILURE THAT PROMPTED THIS TOOL. The ASCO 10007 abstract was
    archived to atlas/data/round227_supplied/ and never given a bibliography entry, so no
    ref-based lookup could reach it, and three rounds later it was requested from the user. A
    document on disk with no ref is invisible to every other check in this repository.
    """
    idx = load()
    bib = yaml.safe_load(open(BIB))["refs"]
    known_names = {re.sub(r"[^a-z0-9]", "", k) for k in bib}
    orphans = []
    for i in idx["items"]:
        if i["ext"] not in {".pdf", ".docx", ".txt", ".xml"}:
            continue
        if i["ref_in_filename"]:
            continue
        base = os.path.basename(i["path"]).lower()
        if base.startswith(("manifest", "search", "round2")) and i["ext"] == ".txt":
            continue          # tool output and manifests are not sources
        stem = re.sub(r"[^a-z0-9]", "", base)
        if any(k in stem for k in known_names):
            continue
        if i["chars"] > 1500:
            orphans.append(i)
    print(f"{len(orphans)} substantial held documents have NO bibliography ref naming them.\n"
          f"Each is a source this atlas can read but cannot cite, and therefore cannot remember:\n")
    for i in sorted(orphans, key=lambda x: -x["chars"])[:40]:
        print(f"  {i['chars']:>8,} chars  {i['path']}")
        print(f"            {i['snippet'][:120]}")
    return 1 if orphans else 0


def cmd_unread():
    bib = yaml.safe_load(open(BIB))["refs"]
    idx = load()
    held = set(idx["held_refs_by_filename"])
    rows = []
    for k, v in bib.items():
        if k in held:
            continue
        if v.get("type") == "primary_abstract_only":
            rows.append((v.get("cited_by", 0) or 0, k, str(v.get("title", ""))[:80]))
    rows.sort(reverse=True)
    print(f"{len(rows)} abstract-only refs with no local document, most-cited first:\n")
    for c, k, t in rows[:40]:
        print(f"  cited_by={c:<4} {k:<24}{t}")
    return 0


def main():
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[1])
    ap.add_argument("--find")
    ap.add_argument("--have")
    ap.add_argument("--check-asks", action="store_true")
    ap.add_argument("--unread", action="store_true")
    ap.add_argument("--orphans", action="store_true")
    a = ap.parse_args()
    if a.find:
        sys.exit(cmd_find(a.find))
    if a.have:
        sys.exit(cmd_have(a.have))
    if a.check_asks:
        sys.exit(cmd_check_asks())
    if a.orphans:
        sys.exit(cmd_orphans())
    if a.unread:
        sys.exit(cmd_unread())
    idx = build()
    print(f"indexed {idx['n_files']} files, {idx['n_bytes']:,} bytes")
    print(f"refs held as a named document: {len(idx['held_refs_by_filename'])}")
    print(f"wrote {os.path.relpath(INDEX_MD, ROOT)} and {os.path.relpath(INDEX_JSON, ROOT)}")


if __name__ == "__main__":
    main()
